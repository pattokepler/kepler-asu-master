import os
import time
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import streamlit as st
from langchain_core.messages import AIMessage, HumanMessage
from langchain_community.vectorstores import FAISS
from sentence_transformers import SentenceTransformer
from langchain_groq.chat_models import ChatGroq
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx
from langchain.embeddings.base import Embeddings

# Create .streamlit directory and config.toml if they don't exist
os.makedirs('.streamlit', exist_ok=True)
with open('.streamlit/config.toml', 'w') as f:
    f.write('[client]\ntoolbarMode = "minimal"')

# Create a custom embeddings class
class SentenceTransformersEmbeddings(Embeddings):
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, texts: list) -> list:
        return self.model.encode(texts, convert_to_tensor=True).tolist()

    def embed_query(self, text: str) -> list:
        return self.model.encode(text, convert_to_tensor=True).tolist()

# Function to handle processing the specific .docx
def process_docx():
    docx_filename = "2025_FAQ.docx"
    
    if not os.path.exists(docx_filename):
        raise ValueError(f".docx file '{docx_filename}' not found in the current directory.")
    
    temp_files = []
    try:
        # Read text from the .docx file
        print(f"Reading .docx file: {docx_filename}")
        doc = docx.Document(docx_filename)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

        if not text:
            raise ValueError(f"No text extracted from .docx '{docx_filename}'. Please check the file content.")
        
        # Create Document objects for each chunk of text
        documents = [Document(page_content=text)]
        
        # Split the extracted text into smaller chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        splits = text_splitter.split_documents(documents)
        
        # Use SentenceTransformers for embeddings (no API key required)
        embeddings = SentenceTransformersEmbeddings(model_name='all-MiniLM-L6-v2')  # Using custom embeddings
        
        # Create a vector store from the splits using Sentence-Transformers embeddings
        vectorstore = FAISS.from_documents(splits, embeddings)
        
        return vectorstore
        
    finally:
        # Cleanup any temporary files (if any were created)
        for temp_file in temp_files:
            safe_remove(temp_file)

def safe_remove(filepath):
    """Safely remove a file with retries."""
    max_attempts = 5
    for attempt in range(max_attempts):
        try:
            if os.path.exists(filepath):
                os.unlink(filepath)
            break
        except PermissionError:
            if attempt < max_attempts - 1:
                time.sleep(0.1)
            else:
                print(f"Warning: Could not remove temporary file {filepath}")

# Function to get the response based on user query
def get_response(user_query, chat_history, vectorstore=None):
    # Initialize ChatGroq model with your API key
    llm = ChatGroq(
        model_name="mixtral-8x7b-32768",
        api_key="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"  # Replace with your actual API key
    )
    
    if vectorstore:
        # Perform similarity search if vectorstore is available
        relevant_docs = vectorstore.similarity_search(user_query, k=3)
        relevant_content = "\n".join([doc.page_content for doc in relevant_docs])
        
        template = """
        You are a helpful assistant. Use the following pieces of context from the uploaded document to help answer the question.
        If the context isn't relevant, just answer based on your knowledge.
        
        Context from document: {context}
        
        Chat history: {chat_history}
        User question: {user_question}
        """
        
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()
        
        return chain.stream({
            "context": relevant_content,
            "chat_history": chat_history,
            "user_question": user_query
        })
    else:
        template = """
        You are a helpful assistant. 
        Answer the following questions considering the history of the conversation:
        Chat history: {chat_history}
        User question: {user_question}
        """
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()
        
        return chain.stream({
            "chat_history": chat_history,
            "user_question": user_query
        })

# Streamlit UI
st.set_page_config(page_title="Kepler ASU Chatbot", page_icon="🤖")
st.image("logo.png", width=800)
st.markdown("### Chatbot Assistant", unsafe_allow_html=True)

# Sidebar info
st.sidebar.image("qr-code.png", width=200)
st.sidebar.markdown("<small>keplerasuscholars@asu.edu </small>", unsafe_allow_html=True)

# Initialize vectorstore if not available
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None

# Process the specific .docx file (2025_FAQ_Frequently_Asked_Questions ASU_Kepler.docx)
if not st.session_state.vectorstore:  # Only process if vectorstore is not already available
    try:
        with st.spinner("Processing document..."):
            st.session_state.vectorstore = process_docx()
        st.sidebar.success("ChatBot Initialized OK!")
    except Exception as e:
        st.sidebar.error(f"Error processing .docx: {str(e)}")

# Initialize chat history if not already initialized
if "chat_history" not in st.session_state:
    st.session_state.chat_history = [
        AIMessage(content="Hi, I'm a bot. How can I help you? Ask questions about this scholar's program, and I'll answer them!")
    ]

# Display chat history with delete buttons
for idx, message in enumerate(st.session_state.chat_history):
    with st.chat_message("AI" if isinstance(message, AIMessage) else "Human"):
        col1, col2 = st.columns([0.9, 0.1])
        with col1:
            st.write(message.content)
        with col2:
            if isinstance(message, HumanMessage):  # Only show delete button for user messages
                if st.button("🗑️", key=f"delete_{idx}", help="Delete this message"):
                    st.session_state.chat_history.pop(idx)  # Delete message from history

# Chat input
user_query = st.chat_input("Type your message here...")

if user_query is not None and user_query != "":
    st.session_state.chat_history.append(HumanMessage(content=user_query))
    
    with st.chat_message("Human"):
        st.markdown(user_query)
    
    with st.chat_message("AI"):
        response = st.write_stream(get_response(
            user_query, 
            st.session_state.chat_history,
            st.session_state.vectorstore
        ))
    
    st.session_state.chat_history.append(AIMessage(content=response))
